# Import necessary modules
import re  # Regular expressions for text processing
import os  # File and directory operations
import fitz  # PDF processing library (PyMuPDF)
import time  # Time management for delays and measuring execution time
import torch  # PyTorch library for deep learning
import whisper  # Speech-to-text library
import keyboard  # Keyboard input listener
import threading  # Multithreading support
from docx import Document  # For handling .docx files
from threading import Thread  # Threading for parallel processing
from pydub import AudioSegment  # Audio file manipulation
from accelerate import Accelerator  # Multi-GPU acceleration
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig, GenerationConfig, TextIteratorStreamer

# Function to remove extra spaces from the text
def remove_extra_spaces(text):
    return re.sub(r' {2,}', ' ', text)

# Function to remove special placeholder phrases from the text
def remove_special_phrases(text):
    text = re.sub(r'\[\s*Inaudible\s*\]|\[\s*Silence\s*\]|\[\s*INAUDIBLE\s*\]|\[\s*inaudible\s*\]|- ', '', text)
    text = re.sub(r' \.', '.', text)  # Fix misplaced periods
    return text

# Function to analyze and identify font styles in a PDF document
def fonts(doc):
    styles = {}
    font_counts = {}
    for page in doc:  # Iterate over pages in the PDF
        blocks = page.get_text("dict")["blocks"]  # Extract text blocks
        for b in blocks:  
            if b['type'] == 0:  # Type 0 indicates text content
                for l in b["lines"]:  
                    for s in l["spans"]:  # Process spans of text
                        identifier = (s['size'], s['flags'])  # Unique identifier for font style
                        styles[identifier] = {'size': s['size'], 'flags': s['flags']}
                        font_counts[identifier] = font_counts.get(identifier, 0) + 1
    # Sort fonts by frequency of occurrence
    font_counts = sorted(font_counts.items(), key=lambda x: x[1], reverse=True)
    return font_counts, styles

# Function to consolidate paragraphs into chunks of manageable length
def consolidate_content(paragraphs):
    consolidated_paragraphs = []
    current_paragraph = [paragraphs[0]]
    
    for paragraph in paragraphs[1:]:
        # Merge paragraphs if the combined word count is within the limit
        if sum(len(p.split()) for p in current_paragraph) + len(paragraph.split()) <= 600:
            current_paragraph.append(paragraph)
        else:
            consolidated_paragraphs.append('"{}"'.format(' '.join(current_paragraph)))
            current_paragraph = [paragraph]
    consolidated_paragraphs.append('"{}"'.format(' '.join(current_paragraph)))

    return '\n\n'.join(consolidated_paragraphs)

# Function to extract text from a PDF and organize it into paragraphs
def pdf_to_text(pdf_path):
    doc = fitz.open(pdf_path)  # Open the PDF file
    font_counts, styles = fonts(doc)  # Identify common font styles
    most_common_font_id, _ = font_counts[0]  # Select the most common font
    most_common_font = styles[most_common_font_id]

    paragraphs = []
    current_paragraph = []

    for page in doc:  # Iterate through pages in the PDF
        blocks = page.get_text("dict")["blocks"]
        for b in blocks: 
            if b['type'] == 0:  # Process text blocks only
                for l in b["lines"]:  
                    for s in l["spans"]:  
                        if s['size'] == most_common_font['size']:  # Match with common font size
                            current_paragraph.append(s['text'])
                            if s['text'].strip().endswith(('.', '!', '?')):  # Check for end of sentence
                                current_paragraph_text = " ".join(current_paragraph).strip()
                                if len(current_paragraph_text.split()) > 1:
                                    paragraphs.append(current_paragraph_text)
                                    current_paragraph = []

    if current_paragraph:  # Add any remaining text as a paragraph
        paragraphs.append(" ".join(current_paragraph).strip())

    # Merge short paragraphs with the next paragraph
    merged_paragraphs = []
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if len(para) < 150 and i+1 < len(paragraphs):  # Short paragraphs
            para += " " + paragraphs[i+1]
            i += 2  
        else:
            i += 1
            merged_paragraphs.append(para)

    return consolidate_content(merged_paragraphs)

# Function to extract text from a .docx file
def docx_to_text(docx_path):
    doc = Document(docx_path)  # Load the .docx file
    return '\n'.join([para.text for para in doc.paragraphs])  # Concatenate paragraphs

# Function to process various file types (PDF, DOCX, TXT)
def process_file(file_path):
    file_name, file_extension = os.path.splitext(file_path)
    content = ''  # Initialize content variable
    if file_extension == '.pdf':
        content = pdf_to_text(file_path)
    elif file_extension == '.docx':
        content = docx_to_text(file_path)
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        print(f"Unsupported file type: {file_extension}")
        return
    
    if not content.strip():  # Check if the file is empty
        print(f"Processing unsuccessful for {file_path}. No text extracted.")
        return
    
    content = remove_extra_spaces(content)  # Clean up extra spaces
    content = remove_special_phrases(content)  # Remove placeholder phrases
    
    # Save the processed content to a new file
    with open(f"{file_name}_processed.txt", 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"Processed {file_path} successfully.")

# Function to split text into smaller parts with a maximum length
def split_text_into_parts(text, max_length):
    sentences = re.split(r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s", text)  # Split by sentence boundaries
    chunks = []
    current_chunk = ""
    word_count = 0

    for sentence in sentences:
        words = len(sentence.split())
        if word_count + words > max_length:  # Check if adding the sentence exceeds the limit
            chunks.append(f'"{current_chunk.strip()}"')  # Wrap the chunk in quotes
            current_chunk = sentence
            word_count = words
        else:
            current_chunk += " " + sentence
            word_count += words

    if current_chunk.strip() != "":
        chunks.append(f'"{current_chunk.strip()}"')  # Add the final chunk

    return '\n\n'.join(chunks)

# Function to create a list of text sections from a .txt file
def create_list_from_txt(file_path):
    # Open the text file in read mode with UTF-8 encoding to ensure proper text handling
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Check if the file content lacks double line breaks ('\n\n'), indicating it might need splitting
    if "\n\n" not in content:
        content = split_text_into_parts(content, 600)  # Split content into smaller parts, each up to 600 words

        # Save the modified content back to the file to ensure it adheres to the expected format
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    # Split the content into sections using double line breaks as delimiters
    sections = content.split("\n\n")
    
    # Define an introductory prompt to prepend to each section for consistent processing
    intro_text = ("You are a professional notetaker who takes detailed but concise notes. "
                  "Your goal is to summarize the text and clearly articulate the main points "
                  "so students can understand the key ideas in the text at a glance. "
                  "Write a detailed but concise summary of the following paragraph: ")

    # Prepend the intro_text to each section to provide context for the summarization task
    sections_with_intro = [intro_text + section for section in sections]

    return sections_with_intro  # Return the modified sections as a list

# Function to save summaries into a Word document
def save_summaries_to_word(summaries, original_file_name, save_path):
    # Ensure the directory for saving the summaries exists; create it if it doesn't
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    # Extract the base name of the file without its extension
    base_name = os.path.splitext(original_file_name)[0]

    # Create a new Word document to store the summaries
    doc = Document()
    doc.add_paragraph(f"Summary of {base_name}:\n")  # Add a header with the base name of the original file

    # Iterate through each summary and clean up unwanted lines before adding it to the document
    for summary in summaries:
        cleaned_summary = remove_unwanted_line(summary)  # Remove redundant or irrelevant lines
        doc.add_paragraph(cleaned_summary)  # Add the cleaned summary to the document

    # Construct the save path for the Word document
    # If the base name is too long, truncate it to 50 characters and add ellipses
    if len(base_name) > 50:
        base_name = base_name[:50] + "..."
    save_file_path = os.path.join(save_path, f"Summary of {base_name}.docx")

    doc.save(save_file_path)  # Save the Word document to the specified path

# Function to remove unwanted lines and redundant sentences from the text
def remove_unwanted_line(text):
    # Define phrases that should be removed if found at the beginning of a line
    phrases = ["Summary", "Please summarize", "Here\'s", "Please provide", "Here is", 
               "Below is", "This summary", "The summary", "Our summary", "The text", "This text"]
    for phrase in phrases:
        # Use a regular expression to match and remove lines starting with these phrases
        pattern = re.compile(r'^' + re.escape(phrase) + r'.*:.*$', re.MULTILINE)
        text = pattern.sub('', text)
    
    # Split the text into sections based on blank lines
    sections = re.split(r'\n\s*\n', text)
    
    # Iterate through each section and remove duplicate sentences
    new_sections = []
    for section in sections:
        # Use regex to extract sentences that start with a capital letter and end with a period
        sentences = re.findall(r'([A-Z].*?\.)', section)
        unique_sentences = []
        for sentence in sentences:
            if sentence not in unique_sentences:  # Check if the sentence is already included
                unique_sentences.append(sentence)
        new_sections.append(' '.join(unique_sentences))  # Combine unique sentences back into a section
    
    # Join the sections back together with double line breaks
    text = '\n\n'.join(new_sections)
    
    return text.strip()  # Return the cleaned text, removing any leading or trailing whitespace

# Function to generate text using the Llama model
def llama_generate(
    model: AutoModelForCausalLM,
    tokenizer: AutoTokenizer,
    prompt: str,
    max_new_tokens: int = 128,
    temperature: int = 1.0,
) -> str:
    # Tokenize the input prompt and move the input IDs to the GPU
    inputs = tokenizer(prompt, return_tensors="pt")
    input_ids = inputs["input_ids"].to("cuda")

    # Initialize a streamer to handle token generation with specific parameters
    streamer = TextIteratorStreamer(tokenizer, skip_prompt=True, skip_special_tokens=True)

    # Set up the generation configuration with parameters for randomness and token sampling
    generation_config = GenerationConfig(
        do_sample=True,  # Enable sampling for non-deterministic output
        temperature=temperature,  # Control randomness in generation
        top_p=0.2,  # Use nucleus sampling with top-p threshold
        top_k=80,  # Use top-k sampling to limit the vocabulary size
    )

    # Generate text asynchronously to avoid blocking the main thread
    with torch.no_grad():
        generation_kwargs = {
            "input_ids": input_ids,
            "generation_config": generation_config,
            "return_dict_in_generate": True,
            "output_scores": True,
            "pad_token_id": tokenizer.eos_token_id,
            "max_new_tokens": max_new_tokens,
            "streamer": streamer
        }

        thread = Thread(
            target=model.generate,
            kwargs=generation_kwargs
        )

        print("\nResponse: ")
        thread.start()  # Start text generation in a separate thread
        broader_summary = ""
        c = 0

        # Stream the generated text as it becomes available
        for text in streamer:
            if c < 2:  # Format the initial output
                text = text.strip()
                c += 1
            print(text, end="", flush=True)  # Print the text incrementally
            broader_summary += text
        thread.join()  # Wait for the thread to complete

    return broader_summary  # Return the full generated response

# Function to listen for the Escape key press and set a global flag
def listen_for_escape():
    global escape_pressed
    while True:
        if keyboard.is_pressed('esc'):  # Check if the 'esc' key is pressed
            escape_pressed = True  # Set the flag to True
            break  # Exit the loop
        time.sleep(0.1)  # Add a small delay to reduce CPU usage

# Initialize global variables
all_lists = []  # To store all text sections from files
names = []  # To store corresponding file names

# Main execution block
if __name__ == "__main__":
    # Initialize the escape flag to False
    escape_pressed = False

    # Start a listener thread to monitor for the Escape key press
    listener_thread = threading.Thread(target=listen_for_escape)
    listener_thread.start()

    # Define folder paths for input and output
    folder_path = os.path.dirname(__file__)  # Current script directory
    audio_input_path = os.path.join(folder_path, "Audio_Inputs")
    audio_output_path = os.path.join(folder_path, "Audio_Outputs")
    text_output_path = os.path.join(folder_path, "Text_Inputs")
    save_path = os.path.join(folder_path, "Note_Ouputs")

    # Ensure the text input folder exists, creating it if necessary
    if not os.path.exists(text_output_path):
        os.makedirs(text_output_path)

    # Convert various audio formats to WAV format
    for filename in os.listdir(audio_input_path):
        if filename.endswith((".m4a", ".mp3", ".mp4", ".wav")):  # Supported audio formats
            audio = AudioSegment.from_file(os.path.join(audio_input_path, filename))
            audio.export(os.path.join(audio_input_path, os.path.splitext(filename)[0] + ".wav"), format="wav")
            print(f"Converted {filename} to WAV format.")

    # Load the Whisper speech-to-text model
    audio_model = whisper.load_model("small.en")

    # Transcribe WAV files and save their text outputs
    for filename in os.listdir(audio_input_path):
        if filename.endswith(".wav"):
            # Skip if the processed file already exists
            processed_filename = os.path.splitext(filename)[0] + "_processed.txt"
            if processed_filename in os.listdir(text_output_path):
                print(f"Processed file for {filename} already exists. Skipping transcription.")
                continue

            # Transcribe audio and save to text input folder
            result = audio_model.transcribe(os.path.join(audio_input_path, filename))
            with open(os.path.join(text_output_path, os.path.splitext(filename)[0] + ".txt"), "w") as text_file:
                text_file.write(result["text"])
                print(f"Transcribed {filename} and saved the result.")

    try:
        # Process text files to clean and standardize content
        for file_name in os.listdir(text_output_path):
            if file_name == '.DS_Store' or file_name.endswith("_processed.txt"):  # Skip system files
                continue
            file_full_path = os.path.join(text_output_path, file_name)
            process_file(file_full_path)

        # Create lists of text sections for further processing
        for file_name in os.listdir(text_output_path):
            if file_name.endswith("_processed.txt"):
                file_full_path = os.path.join(text_output_path, file_name)
                list = create_list_from_txt(file_full_path)  # Generate text sections
                all_lists.append(list)
                print(f"Created List for {file_name}")
                names.append(file_name)
                print(list[0] + "\n")  # Print the first section for preview
                
    except FileNotFoundError:
        print(f"\033[91mThe folder at {text_output_path} was not found.\033[0m")  # Highlight error in red
    except Exception as e:
        print(f"\033[91mAn error occurred: {e}\033[0m")  # Handle unexpected errors

    # Unload the audio model to free memory
    del audio_model

    # Initialize Accelerator for multi-GPU processing
    accelerator = Accelerator()

    # Load the Llama2 summarizer model and tokenizer
    model_name = "austinm2151/Llama2_Summarizer"
    tokenizer = AutoTokenizer.from_pretrained("austinm2151/Llama2_Summarizer")

    # Configuration for 4-bit quantization to optimize memory usage
    quant_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_use_double_quant=True,
        bnb_4bit_quant_type="nf4",
        bnb_4bit_compute_dtype=torch.bfloat16
    )

    # Load the Llama2 model with the quantization config
    model = AutoModelForCausalLM.from_pretrained(
        model_name,
        quantization_config=quant_config,
        device_map="auto",
        trust_remote_code=True
    )
    model.config.use_cache = False  # Disable caching for QLoRA compatibility

    # Use the Accelerator's device for model processing
    device = accelerator.device

    # Iterate through all text sections and generate summaries
    for count, (current_list, real_name) in enumerate(zip(all_lists, names)):
        summary = []  # Store generated summaries
        list_length = len(current_list)  # Total number of sections
        start_time = time.time()  # Start timer for progress estimation

        for index, item in enumerate(current_list):
            print(f"\n\nPrompt:\n{item}")
            print(f"\nSource {count+1} of {len(all_lists)}: {names[count]}")

            if index == 0:
                print(f"Part {index+1} of {list_length}")
                print(f"Progress: 0.00%")
                print(f"Estimated time left: Unknown")
            else:
                print(f"\nItem {index+1} of {list_length}")
                print(f"Progress: {progress_percentage:.2f}%")
                print(f"Estimated Time Left: {estimated_time_left_hour:02.0f}:{estimated_time_left_min:02.0f}:{estimated_time_left_sec:02.0f}")
            
            # Stop processing if the Escape key is pressed
            if escape_pressed:
                print(f"Escape key pressed! Exiting at item {index + 1} of {len(current_list)}")
                save_summaries_to_word(summary, real_name, save_path)  # Save progress
                break  

            # Generate a summary using the model
            response = llama_generate(
                model, 
                tokenizer,
                item, 
                max_new_tokens=500,
                temperature=0.97
            )
            summary.append(response)  # Append the generated summary

            # Calculate and print progress and estimated time left
            elapsed_time = time.time() - start_time
            avg_time_per_item = elapsed_time / (index + 1)
            remaining_items = list_length - (index + 1)
            estimated_time_left_total_seconds = avg_time_per_item * remaining_items
            estimated_time_left_hour = avg_time_per_item * remaining_items // 3600
            estimated_time_left_min = (estimated_time_left_total_seconds % 3600) // 60
            estimated_time_left_sec = avg_time_per_item * remaining_items % 60
            progress_percentage = ((index + 1) / list_length) * 100

        # Save summaries if not interrupted by the Escape key
        if not escape_pressed:
            save_summaries_to_word(summary, real_name, save_path)
        else:
            escape_pressed = False  # Reset the escape flag for the next iteration

    # Wait for the listener thread to finish before exiting
    listener_thread.join()
    print(f"\nProgress: 100.00%")
    print(f"Estimated time left: 00:00:00.00")